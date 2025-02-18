import os
import re
import io
import sys
import shutil
import django
import logging
from docx import Document
from django.conf import settings
from django.core.files import File
from google.oauth2 import service_account
from googleapiclient.discovery import build
from django.db import connection, transaction
from googleapiclient.http import MediaIoBaseDownload


# Setup
def setup_django():
    """Ensures Django settings are configured before using ORM or settings."""
    project_path =  os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    sys.path.append(project_path)
    if not settings.configured:
        os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'hyrostool.settings')
        django.setup()

setup_django()
from scriptdata.models import AdScript
from scriptdata.utils.embeddings import generate_embedding, index


# Logging
LOG_DIR = os.path.join(settings.BASE_DIR, 'logs')
os.makedirs(LOG_DIR, exist_ok=True)
log_file_path = os.path.join(LOG_DIR, 'process_gdrive_adscripts_log.log')
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(log_file_path),
        logging.StreamHandler(sys.stdout)
    ]
)

logger = logging.getLogger(__name__)


# Config
SERVICE_ACCOUNT_FILE = "galvanized-app-445607-e7-1604e087ad17.json"
GDRIVE_FOLDER_ID = os.environ.get("GDRIVE_FOLDER_ID")

SCOPES = ["https://www.googleapis.com/auth/drive", "https://www.googleapis.com/auth/documents.readonly"]
creds = service_account.Credentials.from_service_account_file(SERVICE_ACCOUNT_FILE, scopes=SCOPES)

drive_service = build("drive", "v3", credentials=creds)
docs_service = build("docs", "v1", credentials=creds)
INTEGRATIONS = ('FACEBOOK', 'FB', 'YT', 'YOUTUBE', 'GOOGLE', 'SNAPCHAT', 'TIKTOK', 'TWITTER', 'LINKEDIN')
TEMP_FOLDER = os.path.join(settings.BASE_DIR, "temp")
os.makedirs(TEMP_FOLDER, exist_ok=True)


def extract_metadata_from_filename(filename):
    """
    Extract metadata from the filename.
    """
    metadata = {
        "platform": None,
        "ad_type": None,
        "industry": None,
    }

    for name in INTEGRATIONS:
        if name.lower() in filename.lower():
            metadata["platform"] = name

    # Extract ad type from keywords
    if "UGC" in filename:
        metadata["ad_type"] = "User-Generated Content"
    elif "Expert" in filename:
        metadata["ad_type"] = "Expert Interview"
    elif "Testimonial" in filename:
        metadata["ad_type"] = "Testimonial"
    return metadata


def extract_metadata_from_content(text):
    """
    Extract industry and additional metadata from the script content.
    """
    metadata = {
        "industry": None
    }

    # Common industry keywords
    industry_keywords = {
        "skincare": ["skin", "moisturizer", "wrinkles", "collagen"],
        "fitness": ["workout", "exercise", "protein", "gym"],
        "finance": ["investment", "money", "trading", "credit score"],
        "tech": ["AI", "software", "machine learning"]
    }

    for industry, keywords in industry_keywords.items():
        if any(keyword in text.lower() for keyword in keywords):
            metadata["industry"] = industry
            break
    return metadata


def remove_highlighted_text(doc):
    """
    Removes words/lines that are highlighted in red from the document.
    """
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.highlight_color and run.font.highlight_color == 6:  # 6 = 'RED' in python-docx
                run.text = ""
    return doc

def extract_text_from_docx(file_path):
    """
    Extracts text from a .docx file and returns it as a string.
    """
    doc = Document(file_path)
    doc = remove_highlighted_text(doc)
    text = "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
    return text


@transaction.atomic
def process_docx(file_path):
    """
    Processes a .docx file, extracts metadata, and saves to the database.
    """
    filename = os.path.basename(file_path)
    text_content = extract_text_from_docx(file_path)
    metadata_from_filename = extract_metadata_from_filename(filename)
    metadata_from_content = extract_metadata_from_content(text_content)

    # Combine metadata sources
    platform = metadata_from_filename.get("platform", "") or ""
    ad_type = metadata_from_filename.get("ad_type", "") or ""
    industry = metadata_from_content.get("industry", "") or metadata_from_filename.get("industry", "") or ""
    with open(file_path, 'rb') as f:
        django_file = File(f)
        ad_script = AdScript.objects.get_or_create(
            filename=filename,
            defaults={
                "platform": platform,
                "ad_type": ad_type,
                "industry": industry,
                "content": text_content
            }
        )
        ad_script[0].ad_file.save(os.path.basename(file_path), django_file, save=True)

    # Generate embedding & save to pinecone
    embedding_vector = generate_embedding(text_content)
    index.upsert([(filename, embedding_vector, {"platform": platform, "ad_type": ad_type, "industry": industry})])
    print(f"Saved: {filename} → Pinecone (Platform: {platform}, Ad Type: {ad_type}, Industry: {industry})")


def batch_process_docs(input_folder):
    """
    Processes all .docx files in the given folder and saves them to the database.
    """
    logger.info("Initiating batch processing of ad script docx...")
    for filename in os.listdir(input_folder):
        logger.info(f"Processing Doc: {filename}")
        if filename.endswith(".docx"):
            file_path = os.path.join(input_folder, filename)
            process_docx(file_path)


def list_google_docs(folder_id):
    """Function to get Google Docs from the specific folder"""
    query = f"'{folder_id}' in parents and mimeType='application/vnd.google-apps.document'"
    results = drive_service.files().list(q=query, fields="files(id, name)").execute()
    return results.get('files', [])


def fetch_gdrive_docs(file_id, file_name):
    """Function to download and save Google Docs as DOCX"""
    request = drive_service.files().export_media(fileId=file_id, mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    file_stream = io.BytesIO()
    downloader = MediaIoBaseDownload(file_stream, request)
    done = False
    while not done:
        status, done = downloader.next_chunk()
    
    # Save DOCX file
    output_filename = re.sub(r'[\/:*?"<>|]', '_', f"{file_name}.docx")
    logger.info(f"Processing: {file_name} ({file_id}) --> {output_filename}.docx")
    file_path = os.path.join(TEMP_FOLDER, output_filename)
    with open(file_path, "wb") as f:
        f.write(file_stream.getvalue())


def process_google_ad_scripts(limit=None):
    docs = list_google_docs(GDRIVE_FOLDER_ID)[:limit]
    if not docs:
        logger.info("No Google Docs found in the specified folder.")
    else:
        for doc in docs:
            fetch_gdrive_docs(doc['id'], doc['name'])

        logger.info("Initiating batch processing...")
        batch_process_docs(TEMP_FOLDER)
        if os.path.exists(TEMP_FOLDER):
            shutil.rmtree(TEMP_FOLDER)
        logger.info("Processing Complete!")


if __name__ == '__main__':
    try:
        process_google_ad_scripts(limit=10)
    except Exception as e:
        logger.error(f"Error while processing google ad scripts: ", exc_info=True)
    finally:
        connection.close()


