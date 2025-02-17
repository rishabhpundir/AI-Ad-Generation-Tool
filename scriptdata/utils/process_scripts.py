import os
import django

# Django Setup
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'hyrostool.settings')
django.setup()

from docx import Document
from django.conf import settings
from django.core.files import File
from scriptdata.models import AdScript
from scriptdata.utils.embeddings import generate_embedding, index


INTEGRATIONS = ('FACEBOOK', 'FB', 'YT', 'YOUTUBE', 'GOOGLE', 'SNAPCHAT', 'TIKTOK', 'TWITTER', 'LINKEDIN')


def extract_metadata_from_filename(filename):
    """
    Extract metadata from the filename.
    """
    metadata = {
        "platform": None,
        "ad_type": None,
        "industry": None,
    }

    for name in INTEGRATIONS:
        if name.lower() in filename.lower():
            metadata["platform"] = name

    # Extract ad type from keywords
    if "UGC" in filename:
        metadata["ad_type"] = "User-Generated Content"
    elif "Expert" in filename:
        metadata["ad_type"] = "Expert Interview"
    elif "Testimonial" in filename:
        metadata["ad_type"] = "Testimonial"
    return metadata


def extract_metadata_from_content(text):
    """
    Extract industry and additional metadata from the script content.
    """
    metadata = {
        "industry": None
    }

    # Common industry keywords
    industry_keywords = {
        "skincare": ["skin", "moisturizer", "wrinkles", "collagen"],
        "fitness": ["workout", "exercise", "protein", "gym"],
        "finance": ["investment", "money", "trading", "credit score"],
        "tech": ["AI", "software", "machine learning"]
    }

    for industry, keywords in industry_keywords.items():
        if any(keyword in text.lower() for keyword in keywords):
            metadata["industry"] = industry
            break
    return metadata


def remove_highlighted_text(doc):
    """
    Removes words/lines that are highlighted in red from the document.
    """
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.highlight_color and run.font.highlight_color == 6:  # 6 = 'RED' in python-docx
                run.text = ""
    return doc

def extract_text_from_docx(file_path):
    """
    Extracts text from a .docx file and returns it as a string.
    """
    doc = Document(file_path)
    doc = remove_highlighted_text(doc)
    text = "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
    return text


def process_docx(file_path):
    """
    Processes a .docx file, extracts metadata, and saves to the database.
    """
    filename = os.path.basename(file_path)
    text_content = extract_text_from_docx(file_path)
    metadata_from_filename = extract_metadata_from_filename(filename)
    metadata_from_content = extract_metadata_from_content(text_content)

    # Combine metadata sources
    platform = metadata_from_filename.get("platform", "") or ""
    ad_type = metadata_from_filename.get("ad_type", "") or ""
    industry = metadata_from_content.get("industry", "") or metadata_from_filename.get("industry", "") or ""
    with open(file_path, 'rb') as f:
        django_file = File(f)
        ad_script = AdScript.objects.get_or_create(
            filename=filename,
            defaults={
                "platform": platform,
                "ad_type": ad_type,
                "industry": industry,
                "content": text_content
            }
        )
        ad_script[0].ad_file.save(os.path.basename(file_path), django_file, save=True)

    # Generate embedding & save to pinecone
    embedding_vector = generate_embedding(text_content)
    index.upsert([(filename, embedding_vector, {"platform": platform, "ad_type": ad_type, "industry": industry})])
    print(f"Saved: {filename} → Pinecone (Platform: {platform}, Ad Type: {ad_type}, Industry: {industry})")


def batch_process_docs(input_folder):
    """
    Processes all .docx files in the given folder and saves them to the database.
    """
    for filename in os.listdir(input_folder):
        if filename.endswith(".docx"):
            file_path = os.path.join(input_folder, filename)
            process_docx(file_path)


if __name__ == "__main__":
    ad_scripts_folder = os.path.join(settings.BASE_DIR, "ad_scripts")
    batch_process_docs(ad_scripts_folder)
    print("Processing Complete!")
